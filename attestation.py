# -*- coding: utf-8 -*-
"""
generate_attestations.py

Script de génération d'attestations Word à partir d'un modèle .docx
et d'un fichier CSV de donataires.

- Remplace les champs {NOM}, {MONTANT}, {DATE} du modèle Word fourni.
- Ne modifie jamais le template : un nouveau .docx est créé par donataire.
- Détecte au mieux les colonnes (prénom, nom, organisation, montant, civilité).
- Permet de forcer certaines colonnes via des arguments CLI.
- Gère des montants tels que "55 + 100" en les additionnant.
- Sortie : un .docx par donataire dans un dossier, et un ZIP optionnel.

Exemple d'utilisation :
    python generate_attestations.py \
        --csv "MJ-FAM -Contacts-MAJ-25-AOUT(Donateurs 2022-2023-2024).csv" \
        --template "DONATION-AMIS-JENISCH.docx" \
        --date "Lausanne, le 16 octobre 2025" \
        --outdir "attestations_word" \
        --zip

Dépendances : pandas, python-docx
    pip install pandas python-docx
"""
import argparse
import os
import re
import sys
import zipfile

# Encodage sûr pour Windows / Mac
DEFAULT_FS_ENCODING = sys.getfilesystemencoding() or "utf-8"

# Détection robuste du séparateur et encodage CSV
import pandas as pd
from docx import Document


# -----------------------
# Utilitaires de parsing
# -----------------------
ORG_KEYWORDS = (
    "SA", "SÀRL", "SARL", "Sarl", "Association", "Fondation", "Société", "GmbH", "AG", "Ltd", "Inc"
)
FEMALE_HINTS = (
    "anne", "anna", "elle", "ette", "ine", "ene", "ène",
    "a", "ia", "ya", "na", "ina", "liane", "iane", "line", "rine",
    "otte", "ille", "ise", "yse", "cie", "lie", "rie", "nie", "xie",
    "zia", "cia", "tia", "ria", "aude", "onde", "hilde", "rude", "ude", "iette"
)
# Petit lexique féminin explicite pour de meilleurs résultats (adaptable)
FEMALE_NAMES = {
    "virginia","gudrun","marianne","cécile","catherine","lise","nicole","sylvie",
    "eliane","blandine","monique","geneviève","laurence","liliane","ursula","herta",
    "paulette","françoise","elisabeth","elisa","christiane","cynthia","efinizia",
    "jacqueline","annemarie","myriam","liliana","anne","ramona","béatrice",
    "vivianne","thérèse","heidi","edith","monika","julia","iris","hélène","pauline",
    "marie","paola","beatrice","francoise","therese","monica","isabelle"
}


def read_csv_any(path: str) -> pd.DataFrame:
    """Lit un CSV en devinant encodage et séparateur."""
    encodings = ["utf-8-sig", "utf-8", "cp1252", "latin1"]
    last_exc = None
    for enc in encodings:
        try:
            df = pd.read_csv(path, sep=None, engine="python", encoding=enc, dtype=str)
            return df.fillna("")
        except Exception as e:
            last_exc = e
    raise RuntimeError(f"Impossible de lire le CSV: {path}\nDernière erreur: {last_exc}")


def normalize_spaces(s: str) -> str:
    import re
    return re.sub(r"\\s+", " ", str(s)).strip()


def find_columns(df: pd.DataFrame):
    """Devine les colonnes de prénom / nom / organisation / montant / civilité."""
    cols = {c: c.lower() for c in df.columns}
    def col_like(*keys):
        for k in cols:
            lk = cols[k]
            if any(kw in lk for kw in keys):
                return k
        return None

    first = col_like("prénom", "prenom", "first")
    last = col_like("nom", "lastname", "last")
    org = col_like("organisation", "société", "societe", "raison", "entreprise", "compagnie", "company")
    civ = col_like("civilit", "titre", "title", "civility")
    amt = col_like("montant", "amount", "don", "contribution")

    return first, last, org, civ, amt


def parse_amount(val: str) -> float:
    """
    Extrait un montant d'une chaîne, gère "55 + 100" -> 155.
    Ignore les codes postaux probables (>= 1000 et <= 9999) si mélangés à d'autres nombres.
    """
    import re
    if not val:
        return 0.0
    s = val.replace("\\xa0", " ")
    # Si on a un pattern type "55 + 100 + 40"
    if "+" in s:
        parts = [p.strip() for p in re.split(r"\\s*\\+\\s*", s) if p.strip()]
        total = 0.0
        for p in parts:
            pc = p.replace(" ", "").replace("\\u00A0", "").replace(".-", "").replace(".–", "")
            pc = pc.replace(".", "").replace(",", ".")
            nums = re.findall(r"\\d+(?:\\.\\d+)?", pc)
            if not nums:
                continue
            valf = float(nums[0])
            if 1000 <= valf <= 9999 and len(nums) > 1:
                continue
            total += valf
        return total

    # Sinon, on prend le dernier nombre « pertinent »
    nums = re.findall(r"\\d{1,5}(?:[.,]\\d+)?", s)
    if not nums:
        return 0.0
    # Retire les CP probables (4 chiffres) si d'autres nombres existent
    cleaned = []
    for n in nums:
        n2 = n.replace(".", "").replace(",", ".")
        try:
            f = float(n2)
        except Exception:
            continue
        if 1000 <= f <= 9999 and len(nums) > 1:
            continue
        cleaned.append(f)
    return cleaned[-1] if cleaned else 0.0


def find_amount_in_row(row, amount_col: str | None) -> float:
    if amount_col and amount_col in row:
        return parse_amount(str(row[amount_col]))
    # Essaye colonnes contenant 'montant'…
    for c in row.index:
        lc = c.lower()
        if any(k in lc for k in ["montant", "amount", "don"]):
            v = parse_amount(str(row[c]))
            if v > 0:
                return v
    # Sinon balaye de droite à gauche (probables colonnes fin de tableau)
    for c in reversed(row.index.tolist()):
        v = parse_amount(str(row[c]))
        if v > 0 and v < 50000:
            return v
    return 0.0


def is_organization(name_fields) -> bool:
    blob = " ".join([str(x) for x in name_fields]).strip()
    return any(k in blob for k in ORG_KEYWORDS)


def infer_civility(firstname: str, fallback="Monsieur/Madame") -> str:
    fn = normalize_spaces(firstname).lower()
    # Cas couples "X et Y"
    if " et " in f" {fn} " or " & " in f" {fn} ":
        return "Monsieur et Madame"
    # Heuristique "féminin" par dictionnaire / terminaison
    base = fn.split()[0] if fn else ""
    if base in FEMALE_NAMES:
        return "Madame"
    if any(base.endswith(suf) for suf in FEMALE_HINTS):
        return "Madame"
    return "Monsieur" if base else fallback


def build_display_name(row, cfg) -> str:
    """
    Construit la chaîne {NOM} à injecter, avec civilité.
    Si 'organisation' détectée, on garde tel quel (sans civilité).
    """
    firstname = row.get(cfg.first_col, "") if cfg.first_col else ""
    lastname = row.get(cfg.last_col, "") if cfg.last_col else ""
    orgname = row.get(cfg.org_col, "") if cfg.org_col else ""

    # Nom « brut » potentiel depuis des colonnes combinées
    if not (firstname or lastname or orgname):
        for c in row.index:
            lc = c.lower()
            if any(k in lc for k in ["nom complet", "donataire", "beneficiaire", "raison sociale", "nom"]):
                candidate = normalize_spaces(str(row[c]))
                if candidate:
                    orgname = candidate
                    break

    # Entreprise / organisation
    if orgname and is_organization([orgname]):
        return normalize_spaces(orgname)

    # Couples du type "Hélène et Fabio" avec nom unique
    if firstname and (" et " in firstname or " & " in firstname):
        civ = "Monsieur et Madame"
        disp = normalize_spaces(f"{civ} {firstname} {lastname}".strip())
        return disp

    # Cas général personne physique
    civ = ""
    if cfg.civ_col and row.get(cfg.civ_col, ""):
        civ = normalize_spaces(str(row[cfg.civ_col]))
    else:
        civ = infer_civility(firstname)

    full = normalize_spaces(f"{civ} {firstname} {lastname}").strip()
    if not firstname and lastname:
        full = normalize_spaces(f"{civ} {lastname}")
    if not full:
        full = normalize_spaces(orgname or lastname or firstname or "Donataire")
    return full


def replace_placeholders(doc: Document, mapping):
    """Remplace les placeholders dans paragraphes ET tableaux."""
    def _replace_in_paragraph(p):
        for key, val in mapping.items():
            if key in p.text:
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)

    # Paragraphes
    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)


def slugify(value: str) -> str:
    import re
    value = normalize_spaces(value)
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = value.strip().replace(" ", "_")
    return value[:120]


def run(csv_path, template_path, date_text, outdir, make_zip, first_col, last_col, org_col, civ_col, amount_col):
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"CSV introuvable: {csv_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Modèle .docx introuvable: {template_path}")

    df = read_csv_any(csv_path)

    # Colonnes auto si non fournies
    auto_first, auto_last, auto_org, auto_civ, auto_amt = find_columns(df)
    first_col = first_col or auto_first
    last_col = last_col or auto_last
    org_col = org_col or auto_org
    civ_col = civ_col or auto_civ
    amount_col = amount_col or auto_amt

    os.makedirs(outdir, exist_ok=True)

    generated_files = []
    for _, row in df.iterrows():
        display_name = build_display_name(row, argparse.Namespace(first_col=first_col, last_col=last_col, org_col=org_col, civ_col=civ_col))
        amount = find_amount_in_row(row, amount_col)
        if amount <= 0:
            continue

        # Ouvre un nouveau Document basé sur le template pour CHAQUE donataire
        doc = Document(template_path)
        mapping = {
            "{NOM}": display_name,
            "{MONTANT}": f"{int(amount) if float(amount).is_integer() else amount:.0f}",
            "{DATE}": date_text,
        }
        replace_placeholders(doc, mapping)

        fname = f"Attestation_{slugify(display_name)}.docx"
        out_path = os.path.join(outdir, fname)
        doc.save(out_path)
        generated_files.append(out_path)

    zip_path = None
    if make_zip and generated_files:
        zip_path = os.path.join(outdir, "Attestations_Dons_2025_Word.zip")
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for fp in generated_files:
                zf.write(fp, arcname=os.path.basename(fp))

    return generated_files, zip_path


def main():
    parser = argparse.ArgumentParser(description="Générer des attestations Word depuis un CSV et un template .docx")
    parser.add_argument("--csv", default="MJ-FAM -Contacts-MAJ-25-AOUT(Donateurs 2022-2023-2024).csv", help="Chemin du CSV")
    parser.add_argument("--template", default="DONATION-AMIS-JENISCH.docx", help="Chemin du modèle .docx avec {NOM}, {MONTANT}, {DATE}")
    parser.add_argument("--date", default="Lausanne, le 16 octobre 2025", help="Texte pour {DATE}")
    parser.add_argument("--outdir", default="attestations_word", help="Dossier de sortie pour les .docx générés")
    parser.add_argument("--zip", action="store_true", help="Créer aussi une archive ZIP de sortie")
    # Colonnes optionnelles (si l'auto-détection ne convient pas)
    parser.add_argument("--first-col", dest="first_col", default=None, help="Nom exact de la colonne 'Prénom'")
    parser.add_argument("--last-col", dest="last_col", default=None, help="Nom exact de la colonne 'Nom'")
    parser.add_argument("--org-col", dest="org_col", default=None, help="Nom exact de la colonne 'Organisation'")
    parser.add_argument("--civ-col", dest="civ_col", default=None, help="Nom exact de la colonne 'Civilité'")
    parser.add_argument("--amount-col", dest="amount_col", default=None, help="Nom exact de la colonne 'Montant'")
    args = parser.parse_args()

    generated_files, zip_path = run(
        csv_path=args.csv,
        template_path=args.template,
        date_text=args.date,
        outdir=args.outdir,
        make_zip=args.zip,
        first_col=args.first_col,
        last_col=args.last_col,
        org_col=args.org_col,
        civ_col=args.civ_col,
        amount_col=args.amount_col,
    )

    print(f"Générés: {len(generated_files)} fichiers dans '{args.outdir}'")
    if zip_path:
        print(f"Archive: {zip_path}")


if __name__ == "__main__":
    main()
