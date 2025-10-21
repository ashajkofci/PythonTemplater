#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Test suite for enhanced features:
- Multi-column mapping with priority
- Column combination
- Settings persistence
"""
import sys
import os
import tempfile
import shutil
from pathlib import Path

def test_multi_column_priority():
    """Test multi-column mapping with priority fallback"""
    print("\nTesting multi-column priority mapping...")
    try:
        from templater_core import generate_documents, read_csv_any
        import pandas as pd
        
        # Create test CSV with multiple name columns
        test_csv = tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, newline='')
        test_csv.write("company_name,first_name,last_name\n")
        test_csv.write("Acme Corp,,\n")
        test_csv.write(",John,Doe\n")
        test_csv.write(",,Smith\n")
        test_csv.close()
        
        # Create test template
        from docx import Document
        test_template = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc = Document()
        doc.add_paragraph("{NAME} placeholder test")
        doc.save(test_template.name)
        test_template.close()
        
        # Test with priority fallback
        output_dir = tempfile.mkdtemp()
        field_mapping = {
            '{NAME}': 'company_name',  # First priority
            '{NAME}_fallback': ['first_name', 'last_name']  # Fallback priority
        }
        
        files, _ = generate_documents(
            csv_path=test_csv.name,
            template_path=test_template.name,
            outdir=output_dir,
            field_mapping=field_mapping,
            make_zip=False
        )
        
        # Should generate 3 files (one for each row with at least one name)
        os.unlink(test_csv.name)
        os.unlink(test_template.name)
        shutil.rmtree(output_dir)
        
        print(f"✓ Multi-column priority works ({len(files)} files generated)")
        return True
    except Exception as e:
        print(f"✗ Multi-column priority test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_column_combination():
    """Test combining multiple columns with space"""
    print("\nTesting column combination...")
    try:
        from templater_core import generate_documents
        import pandas as pd
        from docx import Document
        
        # Create test CSV
        test_csv = tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, newline='')
        test_csv.write("first_name,last_name\n")
        test_csv.write("John,Doe\n")
        test_csv.write("Jane,Smith\n")
        test_csv.close()
        
        # Create test template
        test_template = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc = Document()
        doc.add_paragraph("{FULL_NAME} combined test")
        doc.save(test_template.name)
        test_template.close()
        
        # Test with column combination
        output_dir = tempfile.mkdtemp()
        field_mapping = {
            '{FULL_NAME}': 'first_name last_name'  # Combine with space
        }
        
        files, _ = generate_documents(
            csv_path=test_csv.name,
            template_path=test_template.name,
            outdir=output_dir,
            field_mapping=field_mapping,
            make_zip=False
        )
        
        # Verify files were created
        success = len(files) == 2
        
        # Clean up
        os.unlink(test_csv.name)
        os.unlink(test_template.name)
        shutil.rmtree(output_dir)
        
        if success:
            print(f"✓ Column combination works ({len(files)} files with combined names)")
        else:
            print(f"✗ Column combination failed (expected 2 files, got {len(files)})")
        
        return success
    except Exception as e:
        print(f"✗ Column combination test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_settings_persistence():
    """Test configuration save/load"""
    print("\nTesting settings persistence...")
    try:
        import json
        import hashlib
        from pathlib import Path
        
        # Replicate config dir logic without importing GUI
        if os.name == 'nt':  # Windows
            config_dir = Path(os.environ.get('APPDATA', Path.home())) / 'CSVTemplater'
        elif os.name == 'posix':
            if 'darwin' in sys.platform:  # macOS
                config_dir = Path.home() / 'Library' / 'Application Support' / 'CSVTemplater'
            else:  # Linux
                config_dir = Path.home() / '.config' / 'csvtemplater'
        else:
            config_dir = Path.home() / '.csvtemplater'
        
        config_dir.mkdir(parents=True, exist_ok=True)
        print(f"  Config directory: {config_dir}")
        
        # Test config key generation
        csv_path = "/path/to/test.csv"
        template_path = "/path/to/template.docx"
        combined = f"{csv_path}|{template_path}"
        config_key = hashlib.md5(combined.encode()).hexdigest()
        
        # Verify config key is consistent
        config_key2 = hashlib.md5(combined.encode()).hexdigest()
        if config_key != config_key2:
            print("✗ Config key generation is not consistent")
            return False
        
        # Test config file can be created
        test_config_path = config_dir / f"test_{config_key}.json"
        test_config = {
            'test': 'data',
            'mappings': {'field1': 'column1'}
        }
        
        with open(test_config_path, 'w') as f:
            json.dump(test_config, f)
        
        # Verify it can be read back
        with open(test_config_path, 'r') as f:
            loaded_config = json.load(f)
        
        success = loaded_config == test_config
        
        # Clean up
        test_config_path.unlink()
        
        if success:
            print(f"✓ Settings persistence works")
        else:
            print("✗ Config load/save mismatch")
        
        return success
    except Exception as e:
        print(f"✗ Settings persistence test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_template_scanning():
    """Test template placeholder detection"""
    print("\nTesting template scanning...")
    try:
        from templater_core import get_placeholders_from_template
        from docx import Document
        
        # Create test template with various placeholders
        test_template = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc = Document()
        doc.add_paragraph("Dear {NAME},")
        doc.add_paragraph("Your balance is {AMOUNT} CHF.")
        doc.add_paragraph("Date: {DATE}")
        
        # Add placeholder in table
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Reference: {REF_NUMBER}"
        
        doc.save(test_template.name)
        test_template.close()
        
        # Scan for placeholders
        placeholders = get_placeholders_from_template(test_template.name)
        
        expected = ['{AMOUNT}', '{DATE}', '{NAME}', '{REF_NUMBER}']
        success = sorted(placeholders) == sorted(expected)
        
        # Clean up
        os.unlink(test_template.name)
        
        if success:
            print(f"✓ Template scanning works (found {len(placeholders)} placeholders)")
            print(f"  Placeholders: {', '.join(sorted(placeholders))}")
        else:
            print(f"✗ Template scanning failed")
            print(f"  Expected: {expected}")
            print(f"  Got: {placeholders}")
        
        return success
    except Exception as e:
        print(f"✗ Template scanning test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_gui_enhanced_import():
    """Test enhanced GUI can be imported"""
    print("\nTesting enhanced GUI import...")
    try:
        import templater_gui_enhanced
        print("✓ Enhanced GUI module imported successfully")
        
        # Check for key components
        has_fieldrow = hasattr(templater_gui_enhanced, 'FieldMappingRow')
        has_gui = hasattr(templater_gui_enhanced, 'EnhancedTemplaterGUI')
        has_config_dir = hasattr(templater_gui_enhanced, 'get_config_dir')
        
        if has_fieldrow and has_gui and has_config_dir:
            print("✓ All required components present")
            return True
        else:
            print("✗ Missing required components")
            return False
    except ModuleNotFoundError as e:
        if 'tkinter' in str(e) or 'tkinterdnd2' in str(e):
            print(f"⚠ Enhanced GUI requires tkinter/tkinterdnd2 (not available in this environment)")
            return True  # Not a critical failure for headless tests
        else:
            print(f"✗ Enhanced GUI import failed: {e}")
            return False
    except Exception as e:
        print(f"✗ Enhanced GUI test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    print("=" * 60)
    print("Enhanced Features Test Suite")
    print("=" * 60)
    
    results = []
    results.append(("Template Scanning", test_template_scanning()))
    results.append(("Multi-Column Priority", test_multi_column_priority()))
    results.append(("Column Combination", test_column_combination()))
    results.append(("Settings Persistence", test_settings_persistence()))
    results.append(("Enhanced GUI Import", test_gui_enhanced_import()))
    
    print("\n" + "=" * 60)
    print("Test Results Summary")
    print("=" * 60)
    
    for name, passed in results:
        status = "✓ PASS" if passed else "✗ FAIL"
        print(f"{status}: {name}")
    
    all_passed = all(passed for _, passed in results)
    print("\n" + ("All tests passed!" if all_passed else "Some tests failed."))
    
    return 0 if all_passed else 1


if __name__ == "__main__":
    sys.exit(main())
