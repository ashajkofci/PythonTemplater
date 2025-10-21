# -*- coding: utf-8 -*-
"""
templater_core.py

Core functionality for CSV to DOCX template generation.
Extracted from attestation.py to be reusable by both CLI and GUI.
"""
import os
import re
import sys
import zipfile
import pandas as pd
from docx import Document

# Encodage sûr pour Windows / Mac
DEFAULT_FS_ENCODING = sys.getfilesystemencoding() or "utf-8"

# Détection robuste du séparateur et encodage CSV
ORG_KEYWORDS = (
    "SA", "SÀRL", "SARL", "Sarl", "Association", "Fondation", "Société", "GmbH", "AG", "Ltd", "Inc"
)
FEMALE_HINTS = (
    "anne", "anna", "elle", "ette", "ine", "ene", "ène",
    "a", "ia", "ya", "na", "ina", "liane", "iane", "line", "rine",
    "otte", "ille", "ise", "yse", "cie", "lie", "rie", "nie", "xie",
    "zia", "cia", "tia", "ria", "aude", "onde", "hilde", "rude", "ude", "iette"
)
FEMALE_NAMES = {
    "virginia","gudrun","marianne","cécile","catherine","lise","nicole","sylvie",
    "eliane","blandine","monique","geneviève","laurence","liliane","ursula","herta",
    "paulette","françoise","elisabeth","elisa","christiane","cynthia","efinizia",
    "jacqueline","annemarie","myriam","liliana","anne","ramona","béatrice",
    "vivianne","thérèse","heidi","edith","monika","julia","iris","hélène","pauline",
    "marie","paola","beatrice","francoise","therese","monica","isabelle"
}


def read_csv_any(path: str) -> pd.DataFrame:
    """Lit un CSV en devinant encodage et séparateur."""
    encodings = ["utf-8-sig", "utf-8", "cp1252", "latin1"]
    last_exc = None
    for enc in encodings:
        try:
            df = pd.read_csv(path, sep=None, engine="python", encoding=enc, dtype=str)
            return df.fillna("")
        except Exception as e:
            last_exc = e
    raise RuntimeError(f"Impossible de lire le CSV: {path}\nDernière erreur: {last_exc}")


def normalize_spaces(s: str) -> str:
    return re.sub(r"\s+", " ", str(s)).strip()


def find_columns(df: pd.DataFrame):
    """Devine les colonnes de prénom / nom / organisation / montant / civilité."""
    cols = {c: c.lower() for c in df.columns}
    def col_like(*keys):
        for k in cols:
            lk = cols[k]
            if any(kw in lk for kw in keys):
                return k
        return None

    first = col_like("prénom", "prenom", "first", "vorname")
    last = col_like("nom", "lastname", "last", "name")
    org = col_like("organisation", "société", "societe", "raison", "entreprise", "compagnie", "company", "institution")
    civ = col_like("civilit", "titre", "title", "civility")
    # Check for "montant" first, then more general terms
    amt = col_like("montant", "amount")
    if not amt:
        amt = col_like("don", "contribution")

    return first, last, org, civ, amt


def parse_amount(val: str) -> float:
    """
    Extrait un montant d'une chaîne, gère "55 + 100" -> 155.
    Ignore les codes postaux probables (>= 1000 et <= 9999) si mélangés à d'autres nombres.
    """
    if not val:
        return 0.0
    s = val.replace("\xa0", " ")
    # Si on a un pattern type "55 + 100 + 40"
    if "+" in s:
        parts = [p.strip() for p in re.split(r"\s*\+\s*", s) if p.strip()]
        total = 0.0
        for p in parts:
            pc = p.replace(" ", "").replace("\u00A0", "").replace(".-", "").replace(".–", "")
            pc = pc.replace(".", "").replace(",", ".")
            nums = re.findall(r"\d+(?:\.\d+)?", pc)
            if not nums:
                continue
            valf = float(nums[0])
            if 1000 <= valf <= 9999 and len(nums) > 1:
                continue
            total += valf
        return total

    # Sinon, on prend le dernier nombre « pertinent »
    nums = re.findall(r"\d{1,5}(?:[.,]\d+)?", s)
    if not nums:
        return 0.0
    # Retire les CP probables (4 chiffres) si d'autres nombres existent
    cleaned = []
    for n in nums:
        n2 = n.replace(".", "").replace(",", ".")
        try:
            f = float(n2)
        except Exception:
            continue
        if 1000 <= f <= 9999 and len(nums) > 1:
            continue
        cleaned.append(f)
    return cleaned[-1] if cleaned else 0.0


def find_amount_in_row(row, amount_col: str | None) -> float:
    if amount_col and amount_col in row:
        return parse_amount(str(row[amount_col]))
    # Essaye colonnes contenant 'montant'…
    for c in row.index:
        lc = c.lower()
        if any(k in lc for k in ["montant", "amount", "don"]):
            v = parse_amount(str(row[c]))
            if v > 0:
                return v
    # Sinon balaye de droite à gauche (probables colonnes fin de tableau)
    for c in reversed(row.index.tolist()):
        v = parse_amount(str(row[c]))
        if v > 0 and v < 50000:
            return v
    return 0.0


def is_organization(name_fields) -> bool:
    blob = " ".join([str(x) for x in name_fields]).strip()
    return any(k in blob for k in ORG_KEYWORDS)


def infer_civility(firstname: str, fallback="Monsieur/Madame") -> str:
    fn = normalize_spaces(firstname).lower()
    # Cas couples "X et Y"
    if " et " in f" {fn} " or " & " in f" {fn} ":
        return "Monsieur et Madame"
    # Heuristique "féminin" par dictionnaire / terminaison
    base = fn.split()[0] if fn else ""
    if base in FEMALE_NAMES:
        return "Madame"
    if any(base.endswith(suf) for suf in FEMALE_HINTS):
        return "Madame"
    return "Monsieur" if base else fallback


def build_display_name(row, first_col, last_col, org_col, civ_col):
    """
    Construit la chaîne {NOM} à injecter, avec civilité.
    Si 'organisation' détectée, on garde tel quel (sans civilité).
    """
    firstname = row.get(first_col, "") if first_col else ""
    lastname = row.get(last_col, "") if last_col else ""
    orgname = row.get(org_col, "") if org_col else ""

    # Nom « brut » potentiel depuis des colonnes combinées
    if not (firstname or lastname or orgname):
        for c in row.index:
            lc = c.lower()
            if any(k in lc for k in ["nom complet", "donataire", "beneficiaire", "raison sociale", "nom"]):
                candidate = normalize_spaces(str(row[c]))
                if candidate:
                    orgname = candidate
                    break

    # Entreprise / organisation
    if orgname and is_organization([orgname]):
        return normalize_spaces(orgname)

    # Couples du type "Hélène et Fabio" avec nom unique
    if firstname and (" et " in firstname or " & " in firstname):
        civ = "Monsieur et Madame"
        disp = normalize_spaces(f"{civ} {firstname} {lastname}".strip())
        return disp

    # Cas général personne physique
    civ = ""
    if civ_col and row.get(civ_col, ""):
        civ = normalize_spaces(str(row[civ_col]))
    else:
        civ = infer_civility(firstname)

    full = normalize_spaces(f"{civ} {firstname} {lastname}").strip()
    if not firstname and lastname:
        full = normalize_spaces(f"{civ} {lastname}")
    if not full:
        full = normalize_spaces(orgname or lastname or firstname or "Donataire")
    return full


def get_placeholders_from_template(template_path: str) -> list:
    """Extract placeholders from a DOCX template (e.g., {NOM}, {MONTANT})."""
    doc = Document(template_path)
    placeholders = set()
    
    # Check paragraphs
    for p in doc.paragraphs:
        matches = re.findall(r'\{[^}]+\}', p.text)
        placeholders.update(matches)
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = re.findall(r'\{[^}]+\}', p.text)
                    placeholders.update(matches)
    
    return sorted(list(placeholders))


def replace_placeholders(doc: Document, mapping):
    """Remplace les placeholders dans paragraphes ET tableaux, même s'ils sont fragmentés entre plusieurs runs."""
    def _replace_in_paragraph(p):
        for key, val in mapping.items():
            if key in p.text:
                # First try simple replacement in individual runs
                replaced = False
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                        replaced = True
                
                # If not replaced, the placeholder might be split across runs
                if not replaced and key in p.text:
                    # Rebuild the paragraph by merging runs
                    full_text = p.text
                    if key in full_text:
                        # Replace in full text
                        new_text = full_text.replace(key, val)
                        
                        # Clear all runs and create a single run with replaced text
                        # Keep the style of the first run
                        if p.runs:
                            first_run = p.runs[0]
                            # Store style properties
                            font_name = first_run.font.name
                            font_size = first_run.font.size
                            bold = first_run.font.bold
                            italic = first_run.font.italic
                            
                            # Clear all runs
                            for run in p.runs[:]:
                                run.text = ''
                            
                            # Set new text in first run
                            first_run.text = new_text
                            # Restore style
                            if font_name:
                                first_run.font.name = font_name
                            if font_size:
                                first_run.font.size = font_size
                            if bold is not None:
                                first_run.font.bold = bold
                            if italic is not None:
                                first_run.font.italic = italic

    # Paragraphes
    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)


def slugify(value: str) -> str:
    value = normalize_spaces(value)
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = value.strip().replace(" ", "_")
    return value[:120]


def generate_documents(csv_path, template_path, outdir, field_mapping, 
                      filename_field=None, filename_prefix="", filename_suffix="",
                      make_zip=False, progress_callback=None):
    """
    Generate DOCX documents from CSV and template.
    
    Args:
        csv_path: Path to CSV file
        template_path: Path to DOCX template
        outdir: Output directory
        field_mapping: Dict mapping placeholder names to CSV column names or column combinations
                      Supports:
                      - Simple string: single column name
                      - Space-separated string: "col1 col2" combines columns
                      - For fallback columns, use placeholder_fallback key with list
        filename_field: CSV column to use for filename (optional)
        filename_prefix: Prefix for output filenames
        filename_suffix: Suffix for output filenames
        make_zip: Whether to create a ZIP archive
        progress_callback: Callback function(current, total, message)
    
    Returns:
        (generated_files, zip_path)
    """
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"CSV introuvable: {csv_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Modèle .docx introuvable: {template_path}")

    df = read_csv_any(csv_path)
    os.makedirs(outdir, exist_ok=True)

    # Validate that mapped columns exist in CSV
    csv_columns = set(df.columns)
    missing_columns = []
    
    for placeholder, csv_spec in field_mapping.items():
        # Skip fallback keys for now (we'll check them separately)
        if placeholder.endswith('_fallback'):
            continue
        
        # Extract column names from the spec
        if isinstance(csv_spec, str) and csv_spec:
            # The runtime behavior is:
            # 1. If spec contains spaces, try to split and combine multiple columns
            # 2. If that yields no results (columns don't exist), try as single column name
            
            # First, check if the whole spec is a valid column name (handles columns with spaces)
            if csv_spec in csv_columns:
                # Valid single column, no problem
                continue
            
            # If not, check if it's meant to be a column combination
            if ' ' in csv_spec:
                col_names = csv_spec.split()
                # Check if any of the split columns exist
                found_any = any(col_name in csv_columns for col_name in col_names)
                
                if not found_any:
                    # None of the columns exist, this is an error
                    missing_columns.append((placeholder, csv_spec))
                # If some exist, that's OK (combination will work with available columns)
            else:
                # Single column that doesn't exist
                missing_columns.append((placeholder, csv_spec))
    
    # Check fallback columns
    for key, fallback_cols in field_mapping.items():
        if key.endswith('_fallback') and isinstance(fallback_cols, list):
            placeholder = key.replace('_fallback', '')
            for col_name in fallback_cols:
                if col_name and col_name not in csv_columns:
                    missing_columns.append((placeholder, col_name))
    
    # If there are missing columns, provide a helpful warning
    if missing_columns:
        missing_details = []
        for placeholder, col_name in missing_columns[:5]:  # Show first 5
            missing_details.append(f"  - {placeholder} → '{col_name}'")
        
        warning_msg = (
            f"Warning: {len(missing_columns)} mapped column(s) not found in CSV.\n"
            f"Available columns: {', '.join(sorted(csv_columns)[:10])}"
            f"{'...' if len(csv_columns) > 10 else ''}\n\n"
            f"Missing mappings:\n" + "\n".join(missing_details)
        )
        
        if len(missing_columns) > 5:
            warning_msg += f"\n  ... and {len(missing_columns) - 5} more"
        
        # Raise a descriptive error so users know what went wrong
        raise ValueError(warning_msg)

    generated_files = []
    total_rows = len(df)
    skipped_rows = []  # Track skipped rows for debugging
    
    print(f"[DEBUG] Starting to process {total_rows} CSV rows...")
    
    for idx, row in df.iterrows():
        if progress_callback:
            progress_callback(idx, total_rows, f"Processing row {idx+1}/{total_rows}")
        
        # Build mapping from template placeholders to values
        mapping = {}
        skip_row = False
        
        for placeholder, csv_spec in field_mapping.items():
            # Skip fallback keys (processed separately)
            if placeholder.endswith('_fallback'):
                continue
            
            value = ""
            
            # Handle empty csv_spec (unmapped placeholder)
            if not csv_spec or csv_spec == "":
                # Leave value as empty string - this is intentional for unmapped placeholders
                mapping[placeholder] = value
                continue
            
            # Check if it's a single column name (even if it contains spaces)
            if isinstance(csv_spec, str) and csv_spec in row:
                # It's a valid single column name, even if it has spaces
                value = str(row[csv_spec]).strip()
            elif isinstance(csv_spec, str) and ' ' in csv_spec:
                # It might be a combination of columns (space-separated)
                parts = []
                for col_name in csv_spec.split():
                    if col_name in row:
                        col_val = str(row[col_name]).strip()
                        if col_val:
                            parts.append(col_val)
                value = ' '.join(parts)
            
            # Try fallback columns if value is still empty
            if isinstance(csv_spec, str) and not value:
                fallback_key = f"{placeholder}_fallback"
                if fallback_key in field_mapping:
                    fallback_cols = field_mapping[fallback_key]
                    for col in fallback_cols:
                        if col in row:
                            val = str(row[col]).strip()
                            if val:
                                value = val
                                break
            
            mapping[placeholder] = value
        
        # Skip ONLY if the row has no data at all in ANY CSV column
        # This is more robust than checking mapped placeholders, which might be unmapped
        row_has_data = any(str(row[col]).strip() for col in df.columns)
        
        if not row_has_data:
            skipped_rows.append(idx + 1)  # +1 for 1-based row number
            continue
        
        # Open a new Document based on the template for EACH row
        doc = Document(template_path)
        replace_placeholders(doc, mapping)
        
        # Determine filename
        if filename_field:
            # Check if it's a combination of fields (space-separated)
            if ' ' in filename_field:
                parts = []
                for col_name in filename_field.split():
                    # Check if it's a template placeholder
                    if col_name.startswith('__TEMPLATE__'):
                        # Extract placeholder name and use its mapped value
                        placeholder = col_name.replace('__TEMPLATE__', '')
                        if placeholder in mapping:
                            val = str(mapping[placeholder]).strip()
                            if val:
                                parts.append(val)
                    elif col_name in row:
                        # Regular CSV column
                        col_val = str(row[col_name]).strip()
                        if col_val:
                            parts.append(col_val)
                base_name = slugify('_'.join(parts)) if parts else f"document_{idx}"
            elif filename_field.startswith('__TEMPLATE__'):
                # Single template placeholder
                placeholder = filename_field.replace('__TEMPLATE__', '')
                if placeholder in mapping:
                    base_name = slugify(str(mapping[placeholder]))
                else:
                    base_name = f"document_{idx}"
            elif filename_field in row:
                # Single CSV column
                base_name = slugify(str(row[filename_field]))
            else:
                # Use first non-empty value as fallback
                base_name = slugify(next((v for v in mapping.values() if v), f"document_{idx}"))
        else:
            # Use first non-empty value as fallback
            base_name = slugify(next((v for v in mapping.values() if v), f"document_{idx}"))
        
        fname = f"{filename_prefix}{base_name}{filename_suffix}.docx"
        out_path = os.path.join(outdir, fname)
        
        # Handle duplicate filenames
        counter = 1
        while os.path.exists(out_path):
            fname = f"{filename_prefix}{base_name}_{counter}{filename_suffix}.docx"
            out_path = os.path.join(outdir, fname)
            counter += 1
        
        doc.save(out_path)
        generated_files.append(out_path)

    zip_path = None
    if make_zip and generated_files:
        zip_path = os.path.join(outdir, "generated_documents.zip")
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for fp in generated_files:
                zf.write(fp, arcname=os.path.basename(fp))
    
    # Debug output
    print(f"[DEBUG] Document generation summary:")
    print(f"[DEBUG]   - Total CSV rows: {total_rows}")
    print(f"[DEBUG]   - Documents generated: {len(generated_files)}")
    print(f"[DEBUG]   - Rows skipped: {len(skipped_rows)}")
    if skipped_rows and len(skipped_rows) <= 10:
        print(f"[DEBUG]   - Skipped row numbers: {', '.join(map(str, skipped_rows))}")
    elif skipped_rows:
        print(f"[DEBUG]   - Skipped row numbers (first 10): {', '.join(map(str, skipped_rows[:10]))}...")
    
    if progress_callback:
        progress_callback(total_rows, total_rows, f"Complete! Generated {len(generated_files)} files")

    return generated_files, zip_path
